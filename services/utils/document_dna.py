"""
Document DNA - Deep Metadata Extraction

Extraherar omfattande metadata från filer för:
- Deduplicering (SHA256 hash)
- Kvalitetsbedömning (creation_tool)
- Sökbarhet (author, title, sheet_names)
- Spårbarhet (provenance)

Princip: HARDFAIL > Silent Fallback
"""

import os
import stat
import hashlib
from datetime import datetime, timezone
import logging
from zoneinfo import ZoneInfo

LOGGER = logging.getLogger('DocumentDNA')

# --- OPTIONAL DEPENDENCIES ---
# Safe imports with fallback flags

try:
    import magic
    MAGIC_AVAILABLE = True
except ImportError as e:
    MAGIC_AVAILABLE = False
    LOGGER.info(f"python-magic ej tillgängligt: {e}")

try:
    from langdetect import detect, LangDetectException
    LANGDETECT_AVAILABLE = True
except ImportError as e:
    LANGDETECT_AVAILABLE = False
    LOGGER.info(f"langdetect ej tillgängligt: {e}")

try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError as e:
    PIL_AVAILABLE = False
    LOGGER.info(f"Pillow ej tillgängligt: {e}")

try:
    import fitz  # pymupdf
    FITZ_AVAILABLE = True
except ImportError as e:
    FITZ_AVAILABLE = False
    LOGGER.info(f"pymupdf (fitz) ej tillgängligt: {e}")

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError as e:
    DOCX_AVAILABLE = False
    LOGGER.info(f"python-docx ej tillgängligt: {e}")

try:
    import openpyxl
    OPENPYXL_AVAILABLE = True
except ImportError as e:
    OPENPYXL_AVAILABLE = False
    LOGGER.info(f"openpyxl ej tillgängligt: {e}")

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError as e:
    PANDAS_AVAILABLE = False
    LOGGER.info(f"pandas ej tillgängligt: {e}")

# Timezone
try:
    SYSTEM_TZ = ZoneInfo("Europe/Stockholm")
except Exception:
    LOGGER.warning("Kunde inte ladda Europe/Stockholm, använder UTC")
    SYSTEM_TZ = timezone.utc


# =============================================================================
# MAIN ENTRY POINTS
# =============================================================================

def extract_document_dna(filepath: str) -> dict:
    """
    Main entry point. Extraherar full metadata från en fil.
    
    Args:
        filepath: Sökväg till filen
    
    Returns:
        dict med file_dna, intrinsic, content_metrics, provenance
    """
    if not os.path.exists(filepath):
        LOGGER.error(f"HARDFAIL: Fil finns inte: {filepath}")
        return {
            "file_dna": {},
            "intrinsic": {},
            "content_metrics": {},
            "provenance": {},
            "error": f"File not found: {filepath}"
        }
    
    # Extrahera alla delar
    file_dna = _extract_file_dna(filepath)
    mime_type = file_dna.get('mime_type_detected', '')
    
    intrinsic = _extract_intrinsic_properties(filepath, mime_type)
    content_metrics = _extract_content_metrics(filepath, mime_type)
    provenance = _extract_provenance(filepath)
    
    # Lägg till tool_classification för quality scoring
    if intrinsic.get('creation_tool'):
        intrinsic['tool_classification'] = _classify_creation_tool(intrinsic['creation_tool'])
    
    return {
        "file_dna": file_dna,
        "intrinsic": intrinsic,
        "content_metrics": content_metrics,
        "provenance": provenance
    }


def _format_dna_for_llm(dna: dict) -> str:
    """
    Token Economy: Filtrera DNA för LLM-konsumtion.
    
    INCLUDE: title_embedded, author_embedded, creation_tool,
             sheet_names, page_count, language_detected
    
    EXCLUDE: file_hash, file_permissions, file_size_bytes,
             absolute_path, character_count
    
    Returns:
        Formaterad sträng för prompt injection
    """
    parts = []
    
    intrinsic = dna.get('intrinsic', {})
    content = dna.get('content_metrics', {})
    
    # Title
    if intrinsic.get('title_embedded'):
        parts.append(f"TITEL: {intrinsic['title_embedded']}")
    
    # Author
    if intrinsic.get('author_embedded'):
        parts.append(f"FÖRFATTARE: {intrinsic['author_embedded']}")
    
    # Creation tool + classification
    if intrinsic.get('creation_tool'):
        tool = intrinsic['creation_tool']
        classification = _classify_creation_tool(tool)
        parts.append(f"VERKTYG: {tool} ({classification})")
    
    # Page count
    if intrinsic.get('page_count'):
        parts.append(f"SIDOR: {intrinsic['page_count']}")
    
    # Sheet names (Excel)
    if intrinsic.get('sheet_names'):
        sheets = ', '.join(intrinsic['sheet_names'][:5])  # Max 5
        if len(intrinsic['sheet_names']) > 5:
            sheets += f" (+{len(intrinsic['sheet_names']) - 5} till)"
        parts.append(f"FLIKAR: {sheets}")
    
    # Language
    if content.get('language_detected'):
        parts.append(f"SPRÅK: {content['language_detected']}")
    
    if not parts:
        return ""
    
    return "DOKUMENTKONTEXT:\n" + "\n".join(parts)


# =============================================================================
# FILE DNA (Fingerprint)
# =============================================================================

def _extract_file_dna(filepath: str) -> dict:
    """Extrahera fil-DNA: hash, storlek, MIME, kryptering."""
    return {
        "file_hash_sha256": _hash_file_streaming(filepath),
        "file_size_bytes": _get_file_size(filepath),
        "mime_type_detected": _detect_mime_type(filepath),
        "extension_normalized": os.path.splitext(filepath)[1].lower(),
        "is_encrypted": _check_encryption(filepath)
    }


def _hash_file_streaming(filepath: str, chunk_size: int = 8192) -> str:
    """SHA256 utan att ladda hela filen i RAM."""
    sha256 = hashlib.sha256()
    try:
        with open(filepath, 'rb') as f:
            for chunk in iter(lambda: f.read(chunk_size), b''):
                sha256.update(chunk)
        return sha256.hexdigest()
    except Exception as e:
        LOGGER.error(f"Kunde inte hasha {filepath}: {e}")
        return ""


def _get_file_size(filepath: str) -> int:
    """Hämta filstorlek i bytes."""
    try:
        return os.path.getsize(filepath)
    except Exception as e:
        LOGGER.error(f"Kunde inte hämta storlek för {filepath}: {e}")
        return 0


def _detect_mime_type(filepath: str) -> str:
    """Detektera MIME-typ med libmagic (inte extension)."""
    if not MAGIC_AVAILABLE:
        # Fallback till extension-baserad gissning
        ext = os.path.splitext(filepath)[1].lower()
        mime_map = {
            '.pdf': 'application/pdf',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            '.csv': 'text/csv',
            '.txt': 'text/plain',
            '.md': 'text/markdown',
            '.json': 'application/json',
            '.png': 'image/png',
            '.jpg': 'image/jpeg',
            '.jpeg': 'image/jpeg',
        }
        return mime_map.get(ext, 'application/octet-stream')
    
    try:
        mime = magic.Magic(mime=True)
        return mime.from_file(filepath)
    except Exception as e:
        LOGGER.warning(f"MIME-detection misslyckades för {filepath}: {e}")
        return 'application/octet-stream'


def _check_encryption(filepath: str) -> bool:
    """Kontrollera om filen är krypterad/lösenordsskyddad."""
    ext = os.path.splitext(filepath)[1].lower()
    
    # PDF encryption check
    if ext == '.pdf' and FITZ_AVAILABLE:
        try:
            with fitz.open(filepath) as doc:
                return doc.is_encrypted
        except Exception as e:
            LOGGER.debug(f"PDF encryption check misslyckades: {e}")
    
    # Excel encryption check (simplified)
    if ext in ['.xlsx', '.xls'] and OPENPYXL_AVAILABLE:
        try:
            openpyxl.load_workbook(filepath, read_only=True)
            return False  # Om vi kan öppna, är den inte krypterad
        except Exception as e:
            if 'password' in str(e).lower() or 'encrypted' in str(e).lower():
                return True
            LOGGER.debug(f"Excel encryption check misslyckades: {e}")
    
    return False


# =============================================================================
# INTRINSIC PROPERTIES (Embedded Metadata)
# =============================================================================

def _extract_intrinsic_properties(filepath: str, mime_type: str) -> dict:
    """Extrahera inbäddad metadata baserat på filtyp."""
    result = {}
    ext = os.path.splitext(filepath)[1].lower()
    
    # PDF
    if mime_type == 'application/pdf' or ext == '.pdf':
        result.update(_extract_pdf_metadata(filepath))
    
    # Word
    elif 'wordprocessingml' in mime_type or ext == '.docx':
        result.update(_extract_docx_metadata(filepath))
    
    # Excel
    elif 'spreadsheetml' in mime_type or ext in ['.xlsx', '.xls']:
        result.update(_extract_excel_metadata(filepath))
    
    # CSV
    elif mime_type == 'text/csv' or ext == '.csv':
        result.update(_extract_csv_metadata(filepath))
    
    # Images
    elif mime_type.startswith('image/') or ext in ['.png', '.jpg', '.jpeg', '.gif', '.webp']:
        result.update(_extract_image_metadata(filepath))
    
    return result


def _extract_pdf_metadata(filepath: str) -> dict:
    """Extrahera metadata från PDF."""
    result = {}
    
    if not FITZ_AVAILABLE:
        return result
    
    try:
        with fitz.open(filepath) as doc:
            result['page_count'] = doc.page_count
            
            metadata = doc.metadata
            if metadata:
                if metadata.get('author'):
                    result['author_embedded'] = metadata['author']
                if metadata.get('title'):
                    result['title_embedded'] = metadata['title']
                if metadata.get('creator'):
                    result['creation_tool'] = metadata['creator']
                elif metadata.get('producer'):
                    result['creation_tool'] = metadata['producer']
    except Exception as e:
        LOGGER.warning(f"PDF-metadata extraction misslyckades: {e}")
    
    return result


def _extract_docx_metadata(filepath: str) -> dict:
    """Extrahera metadata från Word-dokument."""
    result = {}
    
    if not DOCX_AVAILABLE:
        return result
    
    try:
        doc = docx.Document(filepath)
        props = doc.core_properties
        
        if props.author:
            result['author_embedded'] = props.author
        if props.title:
            result['title_embedded'] = props.title
        
        # Räkna sidor (approximativt via sections)
        result['page_count'] = len(doc.sections)
        
        # Word skapar alltid med Word
        result['creation_tool'] = 'Microsoft Word'
    except Exception as e:
        LOGGER.warning(f"DOCX-metadata extraction misslyckades: {e}")
    
    return result


def _extract_excel_metadata(filepath: str) -> dict:
    """Extrahera metadata från Excel."""
    result = {}
    
    if not OPENPYXL_AVAILABLE:
        return result
    
    try:
        wb = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
        
        result['sheet_names'] = wb.sheetnames
        
        # Räkna totalt antal rader
        total_rows = 0
        for sheet_name in wb.sheetnames[:10]:  # Max 10 sheets för prestanda
            sheet = wb[sheet_name]
            if sheet.max_row:
                total_rows += sheet.max_row
        result['row_count_total'] = total_rows
        
        # Excel skapar alltid med Excel
        result['creation_tool'] = 'Microsoft Excel'
        
        wb.close()
    except Exception as e:
        LOGGER.warning(f"Excel-metadata extraction misslyckades: {e}")
    
    return result


def _extract_csv_metadata(filepath: str) -> dict:
    """Extrahera metadata från CSV."""
    result = {}
    
    if not PANDAS_AVAILABLE:
        # Fallback: räkna rader manuellt
        try:
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                result['row_count_total'] = sum(1 for _ in f)
        except Exception as e:
            LOGGER.warning(f"CSV row count misslyckades: {e}")
        return result
    
    try:
        # Läs bara header för kolumnnamn
        df = pd.read_csv(filepath, nrows=0)
        result['sheet_names'] = list(df.columns)  # Kolumnnamn som "sheet names"
        
        # Räkna rader (effektivt)
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            result['row_count_total'] = sum(1 for _ in f)
    except Exception as e:
        LOGGER.warning(f"CSV-metadata extraction misslyckades: {e}")
    
    return result


def _extract_image_metadata(filepath: str) -> dict:
    """Extrahera metadata från bilder."""
    result = {}
    
    if not PIL_AVAILABLE:
        return result
    
    try:
        with Image.open(filepath) as img:
            result['dimensions'] = f"{img.width}x{img.height}"
            
            # DPI om tillgängligt
            if 'dpi' in img.info:
                dpi = img.info['dpi']
                if isinstance(dpi, tuple):
                    result['resolution_dpi'] = int(dpi[0])
                else:
                    result['resolution_dpi'] = int(dpi)
    except Exception as e:
        LOGGER.warning(f"Image-metadata extraction misslyckades: {e}")
    
    return result


# =============================================================================
# CONTENT METRICS
# =============================================================================

def _extract_content_metrics(filepath: str, mime_type: str) -> dict:
    """Extrahera innehållsmetriker: ordantal, tecken, språk."""
    result = {}
    
    # Extrahera textsample
    text = _extract_text_sample(filepath, mime_type)
    
    if text:
        text_stripped = text.strip()
        result['empty_file'] = len(text_stripped) == 0
        result['word_count'] = len(text_stripped.split())
        result['character_count'] = len(text_stripped)
        result['language_detected'] = _detect_language(text_stripped)
    else:
        result['empty_file'] = True
        result['word_count'] = 0
        result['character_count'] = 0
    
    return result


def _extract_text_sample(filepath: str, mime_type: str, max_chars: int = 10000) -> str:
    """Extrahera textinnehåll för analys (max max_chars tecken)."""
    ext = os.path.splitext(filepath)[1].lower()
    
    # Plaintext
    if mime_type in ['text/plain', 'text/csv', 'text/markdown'] or ext in ['.txt', '.md', '.csv']:
        try:
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read(max_chars)
        except Exception as e:
            LOGGER.debug(f"Text extraction misslyckades (plaintext): {e}")
    
    # PDF
    if (mime_type == 'application/pdf' or ext == '.pdf') and FITZ_AVAILABLE:
        try:
            with fitz.open(filepath) as doc:
                text = ""
                for page in doc[:5]:  # Max 5 sidor
                    text += page.get_text() or ""
                    if len(text) >= max_chars:
                        break
                return text[:max_chars]
        except Exception as e:
            LOGGER.debug(f"Text extraction misslyckades (PDF): {e}")
    
    # DOCX
    if ('wordprocessingml' in mime_type or ext == '.docx') and DOCX_AVAILABLE:
        try:
            doc = docx.Document(filepath)
            text = "\n".join([p.text for p in doc.paragraphs[:50]])  # Max 50 paragrafer
            return text[:max_chars]
        except Exception as e:
            LOGGER.debug(f"Text extraction misslyckades (DOCX): {e}")
    
    return ""  # Inget textinnehåll kunde extraheras


def _detect_language(text: str) -> str:
    """Detektera språk från text."""
    if not LANGDETECT_AVAILABLE:
        LOGGER.debug("Språkdetektering ej tillgänglig (langdetect saknas)")
        return ""
    
    if not text or len(text.strip()) < 20:
        LOGGER.debug("Text för kort för språkdetektering (<20 tecken)")
        return ""
    
    try:
        return detect(text[:1000])  # Analysera max 1000 tecken
    except LangDetectException as e:
        LOGGER.debug(f"Språkdetektering misslyckades (LangDetect): {e}")
        return ""
    except Exception as e:
        LOGGER.debug(f"Språkdetektering misslyckades: {e}")
        return ""


# =============================================================================
# PROVENANCE (Source Tracking)
# =============================================================================

def _extract_provenance(filepath: str) -> dict:
    """Extrahera provenance: filnamn, sökväg, permissions, timestamps."""
    result = {
        "original_filename": os.path.basename(filepath),
        "absolute_path": os.path.abspath(filepath)
    }
    
    try:
        file_stat = os.stat(filepath)
        
        # Permissions (Unix-style)
        result['file_permissions'] = oct(stat.S_IMODE(file_stat.st_mode))
        
        # Timestamps
        mtime = datetime.fromtimestamp(file_stat.st_mtime, tz=SYSTEM_TZ)
        atime = datetime.fromtimestamp(file_stat.st_atime, tz=SYSTEM_TZ)
        
        result['last_modified_ts'] = mtime.isoformat()
        result['last_accessed_ts'] = atime.isoformat()
        
    except Exception as e:
        LOGGER.warning(f"Provenance extraction misslyckades: {e}")
    
    return result


# =============================================================================
# QUALITY CLASSIFICATION
# =============================================================================

def _classify_creation_tool(tool: str) -> str:
    """
    Klassificera verktyg för kvalitets-scoring.
    
    Returns: 'authoring' | 'scanner' | 'distiller' | 'unknown'
    
    authoring: Ursprungliga dokument (hög kvalitet)
    scanner: Inskannade dokument (potentiellt OCR-fel)
    distiller: PDF-konverterare (varierande kvalitet)
    """
    if not tool:
        return 'unknown'
    
    tool_lower = tool.lower()
    
    # Authoring tools (original documents)
    authoring_patterns = [
        'microsoft word', 'microsoft excel', 'microsoft powerpoint',
        'google docs', 'google sheets', 'libreoffice', 'openoffice',
        'pages', 'numbers', 'keynote', 'word', 'excel'
    ]
    if any(p in tool_lower for p in authoring_patterns):
        return 'authoring'
    
    # Scanners
    scanner_patterns = [
        'ricoh', 'canon', 'xerox', 'fujitsu', 'scansnap', 'epson',
        'scanner', 'scan', 'ocr', 'abbyy', 'readiris'
    ]
    if any(p in tool_lower for p in scanner_patterns):
        return 'scanner'
    
    # Distillers/converters
    distiller_patterns = [
        'distiller', 'pdf', 'acrobat', 'ghostscript', 'primopdf',
        'cutepdf', 'bullzip', 'doPDF', 'nitro'
    ]
    if any(p in tool_lower for p in distiller_patterns):
        return 'distiller'
    
    return 'unknown'


# =============================================================================
# EXPORTS
# =============================================================================

__all__ = [
    'extract_document_dna',
    '_format_dna_for_llm',
    '_classify_creation_tool',
    '_hash_file_streaming',
]


# =============================================================================
# TEST
# =============================================================================

if __name__ == "__main__":
    import sys
    
    if len(sys.argv) > 1:
        filepath = sys.argv[1]
        print(f"Extracting DNA from: {filepath}")
        dna = extract_document_dna(filepath)
        
        import json
        print(json.dumps(dna, indent=2, ensure_ascii=False, default=str))
        
        print("\n--- LLM Format ---")
        print(_format_dna_for_llm(dna))
    else:
        print("Usage: python document_dna.py <filepath>")

